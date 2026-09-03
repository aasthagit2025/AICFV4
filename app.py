"""Analytical Insight Confidence Framework - Streamlit application (v6).

Structural changes versus v5:

* Nothing is computed until the researcher presses "Run AICF validation".
  v5 re-parsed every uploaded file and re-scored every insight on every widget
  interaction, including each keystroke in the objective box. With large
  PowerPoint and Excel uploads that was the main source of slowness and timeouts
  on Streamlit Community Cloud.
* All file parsing is cached on the file's content hash, so a re-run after
  changing one dropdown reuses the parsed data.
* Results are held in session state and shown in tabs instead of one wide table.
* A Diagnostics tab reports what the parser actually found in each file, so a run
  that produces no cross-source validation can be explained rather than guessed at.
"""

from __future__ import annotations

import hashlib
import io
import re
import traceback

import pandas as pd
import streamlit as st

from aicf_config import (
    ANALYTICAL_TECHNIQUES,
    APP_VERSION,
    DATA_ENVIRONMENTS,
    DIMENSIONS,
    INDUSTRIES,
    INSIGHT_CUES,
    INSIGHT_SOURCES,
    NON_INSIGHT_CUES,
    NON_INSIGHT_EXACT,
    REQUIRED_COLUMNS,
    STUDY_TYPES,
    canonical,
    canonical_list,
)
from aicf_context_gate import alignment_penalty, context_keywords, count_hits, objective_terms, run_gate
from aicf_evidence import (
    ParseDiagnostics,
    compact_text,
    extract_metric_records,
    file_bytes,
    tokenize,
    validate_insight_against_records,
)
from aicf_exports import build_docx_summary, build_excel_report
from aicf_framework import DIMENSIONS as FRAMEWORK_DIMENSIONS, score_insight, validate_columns
from aicf_qa import add_qa_columns
from analysis_validators import (
    ANALYSIS_VALIDATION_COLUMNS,
    extract_analysis_evidence,
    validate_insight_with_analysis_evidence,
)
from chart_qa_agent import CHART_QA_COLUMNS, extract_ppt_chart_records, validate_chart_claim_against_records
from insight_generator import extract_questionnaire_text, read_survey_file
from ppt_exporter import build_pptx_report
from ppt_insight_reader import read_pptx_insights

st.set_page_config(page_title="AICF Tool", page_icon="AI", layout="wide")


# ---------------------------------------------------------------------------
# Cached file handling
# ---------------------------------------------------------------------------

def _digest(uploaded_file) -> str:
    return hashlib.sha256(file_bytes(uploaded_file)).hexdigest()


class CachedFile(io.BytesIO):
    """A picklable stand-in for a Streamlit UploadedFile, safe inside caches."""

    def __init__(self, name: str, data: bytes):
        super().__init__(data)
        self.name = name

    def clone(self) -> "CachedFile":
        return CachedFile(self.name, self.getvalue())


def to_cached(uploaded_file) -> CachedFile | None:
    if uploaded_file is None:
        return None
    return CachedFile(uploaded_file.name, file_bytes(uploaded_file))


@st.cache_data(show_spinner=False)
def cached_questionnaire_text(name: str, data: bytes) -> str:
    try:
        return extract_questionnaire_text(CachedFile(name, data))
    except Exception as exc:
        return f"[Questionnaire could not be read: {exc}]"


@st.cache_data(show_spinner=False)
def cached_tabular(name: str, data: bytes) -> pd.DataFrame:
    lowered = name.lower()
    handle = CachedFile(name, data)
    if lowered.endswith(".csv"):
        return pd.read_csv(handle)
    if lowered.endswith((".xlsx", ".xlsm", ".xls")):
        return pd.read_excel(handle)
    if lowered.endswith(".sav"):
        frame, _ = read_survey_file(handle)
        return frame
    raise ValueError(f"{name}: upload CSV, Excel or SPSS .sav data.")


def extract_docx_text(handle) -> str:
    from docx import Document

    handle.seek(0)
    document = Document(handle)
    texts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))
    return "\n".join(texts)


def read_docx_insights(name: str, data: bytes) -> pd.DataFrame:
    text = extract_docx_text(CachedFile(name, data))
    rows = []
    for sentence in re.split(r"(?<=[.!?])\s+|\n+", text):
        cleaned = compact_text(sentence, 500)
        words = len(re.findall(r"\w+", cleaned))
        if not 7 <= words <= 90:
            continue
        lowered = cleaned.lower().strip()
        if lowered in NON_INSIGHT_EXACT or any(cue in lowered for cue in NON_INSIGHT_CUES):
            continue
        if count_hits(cleaned, INSIGHT_CUES) == 0:
            continue
        rows.append({
            "insight_id": f"DOCX-{len(rows) + 1:03d}",
            "theme": "Word report insight",
            "insight_text": cleaned,
            "evidence_note": "",
            "source_file": name,
        })
    return pd.DataFrame(rows)


@st.cache_data(show_spinner=False)
def cached_insight_report(name: str, data: bytes) -> pd.DataFrame:
    lowered = name.lower()
    if lowered.endswith((".csv", ".xlsx", ".xlsm", ".xls", ".sav")):
        frame = cached_tabular(name, data)
        frame.columns = [str(column).strip() for column in frame.columns]
        return frame
    if lowered.endswith(".pptx"):
        return read_pptx_insights(CachedFile(name, data))
    if lowered.endswith(".docx"):
        return read_docx_insights(name, data)
    raise ValueError("Upload the insight report as CSV, Excel, SPSS .sav, PowerPoint .pptx or Word .docx.")


# ---------------------------------------------------------------------------
# Context alignment - per insight, not per study
# ---------------------------------------------------------------------------

def add_context_alignment(df: pd.DataFrame, *, objective: str, industry: str, study_type: str,
                          techniques: list[str], qre_text: str, data_df: pd.DataFrame,
                          advanced_text: str, penalty: int) -> pd.DataFrame:
    aligned = df.copy()
    keywords = context_keywords(study_type, industry, techniques)
    terms = objective_terms(objective)
    support_tokens = tokenize(f"{qre_text} {advanced_text} {' '.join(map(str, data_df.columns))}")

    scores, notes = [], []
    for _, row in aligned.iterrows():
        row_text = " ".join(
            str(row.get(column, "") or "")
            for column in ["theme", "insight_text", "evidence_note"]
            if column in aligned.columns
        )
        row_tokens = tokenize(row_text)
        score = 1
        if count_hits(row_text, keywords) >= 2:
            score += 2
        elif count_hits(row_text, keywords) == 1:
            score += 1
        if len(terms & row_tokens) >= 2:
            score += 1
        if len(row_tokens & support_tokens) >= 4:
            score += 1
        score = max(1, min(5, score - penalty))
        scores.append(score)
        if score >= 4:
            notes.append(f"Strong alignment with {study_type}, {industry} and the uploaded study material.")
        elif score == 3:
            notes.append(f"Moderate alignment with {study_type}; verify the source evidence for this claim.")
        else:
            notes.append(
                f"Weak alignment with {study_type}, {industry}, the selected technique(s) or the uploaded evidence."
            )
    aligned["context_alignment_score"] = scores
    aligned["context_alignment_note"] = notes
    return aligned


# ---------------------------------------------------------------------------
# Validation pipeline
# ---------------------------------------------------------------------------

def add_cross_source_validation(df: pd.DataFrame, evidence_files: list[CachedFile],
                                diagnostics: ParseDiagnostics) -> pd.DataFrame:
    enriched = df.copy()
    records = extract_metric_records([f.clone() for f in evidence_files], diagnostics)
    chart_records = extract_ppt_chart_records([f.clone() for f in evidence_files])
    analysis_records = extract_analysis_evidence([f.clone() for f in evidence_files])

    diagnostics.add("(all evidence files)", "-", f"{len(records)} comparable metric row(s) available for matching")
    diagnostics.add("(all evidence files)", "-", f"{len(chart_records)} PowerPoint chart/table record(s)")
    diagnostics.add("(all evidence files)", "-", f"{len(analysis_records)} advanced-analysis record(s)")

    metrics_out, validation_out = [], []
    chart_out = {column: [] for column in CHART_QA_COLUMNS}
    analysis_out = {column: [] for column in ANALYSIS_VALIDATION_COLUMNS}

    for _, row in enriched.iterrows():
        row_dict = row.to_dict()
        metrics, verdict = validate_insight_against_records(row_dict, records)

        chart_result = validate_chart_claim_against_records(row_dict, chart_records)
        chart_status = str(chart_result.get("chart_validation_status", "") or "").upper()
        if chart_status == "FAIL":
            metrics = chart_result.get("quantitative_metrics_validated") or metrics
            verdict = chart_result.get("cross_source_validation") or verdict
        elif chart_status == "PASS" and "no directly matching" in verdict.lower():
            metrics = chart_result.get("quantitative_metrics_validated") or metrics
            verdict = chart_result.get("cross_source_validation") or verdict

        analysis_result = validate_insight_with_analysis_evidence(row_dict, analysis_records)
        analysis_status = str(analysis_result.get("analysis_validation_status", "") or "").upper()
        if analysis_status in {"PASS", "FAIL"}:
            detail = analysis_result.get("analysis_validation_reason", "")
            if detail:
                verdict = compact_text(f"{verdict} | Analysis check: {detail}", 700)

        metrics_out.append(compact_text(metrics, 900))
        validation_out.append(compact_text(verdict, 700))
        for column in CHART_QA_COLUMNS:
            chart_out[column].append(compact_text(chart_result.get(column, ""), 700))
        for column in ANALYSIS_VALIDATION_COLUMNS:
            analysis_out[column].append(compact_text(analysis_result.get(column, ""), 700))

    enriched["quantitative_metrics_validated"] = metrics_out
    enriched["cross_source_validation"] = validation_out
    for column, values in {**chart_out, **analysis_out}.items():
        enriched[column] = values
    return enriched


PRESERVED_COLUMNS = ["source_file", "slide_number"] + CHART_QA_COLUMNS + ANALYSIS_VALIDATION_COLUMNS


def score_dataframe(df: pd.DataFrame, use_manual_scores: bool) -> tuple[pd.DataFrame, list[str]]:
    results, errors = [], []
    for position, (_, row) in enumerate(df.iterrows(), start=1):
        row_dict = row.to_dict()
        try:
            scored = score_insight(row_dict, use_manual_scores=use_manual_scores).to_dict()
        except Exception as exc:  # one bad row must not kill the whole run
            errors.append(f"Row {position} ({row_dict.get('insight_id', 'no id')}): {exc}")
            continue
        for column in PRESERVED_COLUMNS:
            if column in row_dict:
                scored[column] = row_dict[column]
        results.append(scored)
    return pd.DataFrame(results), errors


def run_pipeline(*, qre_file, data_file, report_file, evidence_files, settings, confirmed, use_manual_scores):
    diagnostics = ParseDiagnostics()

    qre_text = cached_questionnaire_text(qre_file.name, qre_file.getvalue()) if qre_file else ""
    diagnostics.add(qre_file.name if qre_file else "-", "-",
                    f"{len(qre_text.split())} words read from questionnaire" if qre_text else "No questionnaire text")

    data_df = pd.DataFrame()
    if data_file:
        try:
            data_df = cached_tabular(data_file.name, data_file.getvalue())
            diagnostics.add(data_file.name, "-", f"{len(data_df)} rows x {len(data_df.columns)} variables")
        except Exception as exc:
            diagnostics.add(data_file.name, "-", "Primary data could not be read", str(exc))

    insight_df = cached_insight_report(report_file.name, report_file.getvalue())
    if "source_file" not in insight_df.columns:
        insight_df["source_file"] = report_file.name
    diagnostics.add(report_file.name, "-", f"{len(insight_df)} insight row(s) extracted")

    missing = validate_columns(list(insight_df.columns))
    if missing:
        return {"error": f"The insight report is missing required column(s): {', '.join(missing)}. "
                         f"Minimum required columns are: {', '.join(REQUIRED_COLUMNS)}."}

    advanced_parts = []
    for handle in evidence_files or []:
        if handle.name.lower().endswith(".pptx"):
            advanced_parts.append(f"{handle.name}: PowerPoint analysis output")
            continue
        try:
            columns = cached_tabular(handle.name, handle.getvalue()).columns[:15]
            advanced_parts.append(f"{handle.name}: {', '.join(map(str, columns))}")
        except Exception as exc:
            diagnostics.add(handle.name, "-", "Could not summarise this analysis file", str(exc))
    advanced_text = " ".join(advanced_parts)

    gate = run_gate(
        objective=settings["study_objective"], industry=settings["industry"],
        technology=settings["technology"], study_type=settings["study_type"],
        techniques=canonical_list(settings["analytical_technique"].split(", ")) if settings["analytical_technique"] else [],
        qre_text=qre_text, data_df=data_df, insight_df=insight_df,
        evidence_present=bool(data_file or evidence_files), additional_context=advanced_text,
    )
    if gate.blocked:
        return {"gate": gate, "error": None, "blocked": True}

    penalty = alignment_penalty(gate, confirmed)
    techniques = canonical_list(settings["analytical_technique"].split(", ")) if settings["analytical_technique"] else []

    aligned = add_context_alignment(
        insight_df, objective=settings["study_objective"], industry=settings["industry"],
        study_type=settings["study_type"], techniques=techniques, qre_text=qre_text,
        data_df=data_df, advanced_text=advanced_text, penalty=penalty,
    )

    # The insight report is also searched for evidence (a deck often carries both
    # the claims and the charts). Deduplicate by content so a file supplied twice
    # is not parsed twice and its metrics are not double-counted.
    candidates = ([data_file] if data_file else []) + list(evidence_files) + [report_file]
    seen: set[str] = set()
    all_evidence = []
    for handle in candidates:
        key = hashlib.sha256(handle.getvalue()).hexdigest()
        if key in seen:
            diagnostics.add(handle.name, "-", "Skipped: identical to a file already parsed")
            continue
        seen.add(key)
        all_evidence.append(handle)
    validated = add_cross_source_validation(aligned, all_evidence, diagnostics)

    for key, value in settings.items():
        if key not in validated.columns:
            validated[key] = value
        else:
            validated[key] = validated[key].fillna("").astype(str)
            validated.loc[validated[key].str.strip().eq(""), key] = value

    report, row_errors = score_dataframe(validated, use_manual_scores)
    if report.empty:
        return {"error": "No insight could be scored. " + (" ".join(row_errors) if row_errors else "")}

    report = add_qa_columns(report)
    return {
        "report": report, "extracted": insight_df, "gate": gate,
        "diagnostics": diagnostics.to_frame(), "row_errors": row_errors,
        "settings": settings, "penalty": penalty, "error": None, "blocked": False,
    }


# ---------------------------------------------------------------------------
# Templates
# ---------------------------------------------------------------------------

def insight_template() -> pd.DataFrame:
    return pd.DataFrame([
        {"insight_id": "I001", "theme": "Overall satisfaction",
         "insight_text": "Satisfaction is moderate-positive, but the 18.4% low-rating share shows improvement is still required.",
         "evidence_note": "Q5 overall satisfaction: mean 3.59/5, top-two-box 59.7%, low ratings 18.4%, n=347."},
        {"insight_id": "I002", "theme": "Overclaim example",
         "insight_text": "All customers are fully satisfied, so no improvement is required.",
         "evidence_note": "Included to show how the framework challenges absolute wording."},
    ])


def advanced_template() -> pd.DataFrame:
    return pd.DataFrame([
        {"analysis_type": "KDA / Driver Analysis",
         "required_columns": "dependent_variable; driver_variable; importance_score_or_coefficient; performance_score; significance_or_p_value; sample_size"},
        {"analysis_type": "Segmentation",
         "required_columns": "segment_name; segment_size; segment_profile; key_differentiating_variables; needs_attitudes_behaviors; segment_attractiveness; business_action"},
        {"analysis_type": "Conjoint",
         "required_columns": "attribute; level_or_item; utility_or_score; importance_or_rank; preference_share_or_selected_pct; sample_size"},
        {"analysis_type": "MaxDiff / TURF",
         "required_columns": "level_or_item; utility_or_score; importance_or_rank; reach; sample_size"},
        {"analysis_type": "Brand Funnel",
         "required_columns": "level_or_item; total_score; banner_or_segment_cut; sample_size"},
        {"analysis_type": "Correspondence Analysis",
         "required_columns": "brand_or_attribute; dimension_1; dimension_2; inertia_or_variance_explained"},
    ])


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

st.title("Analytical Insight Confidence Framework")
st.caption(
    f"{APP_VERSION}. Validate Human, AI or Human+AI insights against the questionnaire, the data and the "
    "analysis output, and produce an auditable Insight Confidence Index with named researcher sign-off."
)

with st.sidebar:
    st.header("Project governance")
    study_objective = st.text_area(
        "Research objective for this study", height=100,
        placeholder="e.g. Identify which of the five tested propositions drives strongest purchase intent among category buyers.",
        help="Write a study-specific objective. It is carried into the AICF output and used for context alignment.",
    )
    industry = st.selectbox("Industry / sector", [""] + INDUSTRIES)
    technology = st.selectbox("Technology / data environment", [""] + DATA_ENVIRONMENTS)
    study_type = st.selectbox("Study type", [""] + STUDY_TYPES)
    analytical_techniques = st.multiselect("Analytical technique", ANALYTICAL_TECHNIQUES)
    insight_source = st.selectbox("Insight source", INSIGHT_SOURCES)
    researcher_owner = st.text_input("Named researcher owner")
    independent_ai_checker = st.text_input("Independent AI checker", value="Second AI model / reviewer")
    benchmark_reference = st.text_input("Benchmark reference", value="Senior researcher benchmark")

    st.header("AICF dimensions")
    for item in DIMENSIONS.values():
        st.write(f"**{item['label']}** - {float(item['weight']):.0%}")

    st.download_button("Download insight CSV template",
                       insight_template().to_csv(index=False).encode("utf-8"),
                       "aicf_input_template.csv", "text/csv")
    st.download_button("Download advanced-analysis templates",
                       advanced_template().to_csv(index=False).encode("utf-8"),
                       "aicf_advanced_analysis_templates.csv", "text/csv")

st.subheader("Upload study materials")

left, right = st.columns(2)
with left:
    questionnaire_file = st.file_uploader("Required: questionnaire / QRE", type=["docx", "txt"])
    data_file = st.file_uploader("Evidence option 1: primary data", type=["csv", "xlsx", "xls", "sav"])
with right:
    insight_report_file = st.file_uploader("Required: insight report", type=["csv", "xlsx", "xls", "pptx", "docx"])
    additional_analysis_files = st.file_uploader(
        "Evidence option 2: analysis / table output", type=["csv", "xlsx", "xls", "sav", "pptx"],
        accept_multiple_files=True,
        help="Banner tables, KDA, MaxDiff, Conjoint, Segmentation, Brand Funnel or other method output.",
    )

option_a, option_b = st.columns(2)
with option_a:
    use_manual_scores = st.checkbox(
        "Use manual evaluator score columns if present", value=False,
        help="Off by default. When on, rows carrying all seven 1-5 score columns use those instead of automatic scoring.",
    )
with option_b:
    confirmed_same_study = st.checkbox(
        "I confirm these files all belong to the same study", value=False,
        help="Advisory context warnings do not stop the run. Leaving this unticked lowers context alignment "
             "and therefore confidence, rather than refusing to score.",
    )

run_clicked = st.button("Run AICF validation", type="primary", width="stretch")

if run_clicked:
    missing = []
    if not questionnaire_file:
        missing.append("questionnaire / QRE")
    if not insight_report_file:
        missing.append("insight report")
    if not (data_file or additional_analysis_files):
        missing.append("at least one evidence source (primary data or analysis / table output)")
    if missing:
        st.error("Cannot run yet. Please upload: " + "; ".join(missing) + ".")
    else:
        settings = {
            "study_objective": study_objective.strip(),
            "industry": industry,
            "technology": technology,
            "study_type": canonical(study_type),
            "analytical_technique": ", ".join(canonical_list(analytical_techniques)),
            "insight_source": insight_source,
            "researcher_owner": researcher_owner.strip(),
            "independent_ai_checker": independent_ai_checker.strip(),
            "benchmark_reference": benchmark_reference.strip(),
        }
        with st.spinner("Reading files, matching evidence and scoring insights..."):
            try:
                st.session_state["aicf_result"] = run_pipeline(
                    qre_file=to_cached(questionnaire_file),
                    data_file=to_cached(data_file),
                    report_file=to_cached(insight_report_file),
                    evidence_files=[to_cached(f) for f in (additional_analysis_files or [])],
                    settings=settings, confirmed=confirmed_same_study,
                    use_manual_scores=use_manual_scores,
                )
            except Exception as exc:
                st.session_state["aicf_result"] = {"error": str(exc), "traceback": traceback.format_exc()}

result = st.session_state.get("aicf_result")

if not result:
    st.info("Upload the questionnaire, the insight report and at least one evidence source, then press "
            "**Run AICF validation**. Nothing is scored until you do.")
    st.dataframe(insight_template(), width="stretch")

elif result.get("error"):
    st.error(result["error"])
    if result.get("traceback"):
        with st.expander("Technical detail"):
            st.code(result["traceback"])

elif result.get("blocked"):
    st.error("AICF cannot score this run yet.")
    for blocker in result["gate"].blockers:
        st.write(f"- {blocker}")

else:
    report = result["report"]
    gate = result["gate"]

    for note in gate.notes:
        st.info(note)
    if gate.advisories:
        with st.expander(f"{len(gate.advisories)} context advisory(ies) - review before relying on these scores",
                         expanded=not confirmed_same_study):
            for advisory in gate.advisories:
                st.write(f"- {advisory}")
            if result["penalty"]:
                st.warning(f"Context alignment reduced by {result['penalty']} point(s) because the advisories were "
                           "not confirmed. Tick the confirmation box above and re-run if the files do belong together.")
    if result["row_errors"]:
        with st.expander(f"{len(result['row_errors'])} row(s) could not be scored"):
            for error in result["row_errors"]:
                st.write(f"- {error}")

    counts = report["ici_classification"].value_counts()
    summary_cols = st.columns(5)
    summary_cols[0].metric("Insights scored", len(report))
    summary_cols[1].metric("High confidence", int(counts.get("High Confidence", 0)))
    summary_cols[2].metric("Medium confidence", int(counts.get("Medium Confidence", 0)))
    summary_cols[3].metric("Low confidence", int(counts.get("Low Confidence", 0)))
    summary_cols[4].metric("Not scorable", int(counts.get("Not Scorable", 0)))

    verified = int((report["evidence_tier"] >= 3).sum())
    st.caption(
        f"{verified} of {len(report)} insight(s) had at least one figure matched to an uploaded source. "
        "Only those can reach High Confidence."
    )

    tabs = st.tabs(["Summary", "Insights", "First-layer QA", "Evidence trail", "Diagnostics", "Downloads"])

    with tabs[0]:
        st.subheader("Confidence distribution")
        st.bar_chart(report["ici_classification"].value_counts())
        st.subheader("Evidence tier distribution")
        st.bar_chart(report["evidence_tier_label"].value_counts())
        flagged = report[~report["quality_challenge_flags"].str.contains("No major quality challenge", na=False)]
        if not flagged.empty:
            st.error(f"{len(flagged)} insight(s) were challenged by the validation layer and need source "
                     "verification before client use.")
            st.dataframe(flagged[["insight_id", "ici_classification", "quality_challenge_flags"]],
                         width="stretch", hide_index=True)
        qa_fails = int((report["qa_overall_status"] == "FAIL").sum())
        if qa_fails:
            st.warning(f"First-layer research QA failed for {qa_fails} insight(s). See the First-layer QA tab.")

    with tabs[1]:
        chosen = st.multiselect("Filter by confidence", list(counts.index), default=list(counts.index))
        view = report[report["ici_classification"].isin(chosen)]
        for _, row in view.iterrows():
            header = f"{row['insight_id']} - {row['ici_classification']} ({row['ici_score']}/100)"
            with st.expander(header):
                st.write(row["insight_text"])
                score_cols = st.columns(7)
                for col, (key, item) in zip(score_cols, FRAMEWORK_DIMENSIONS.items()):
                    col.metric(str(item["label"]).split(" /")[0], f"{row[key]}/5")
                st.write(f"**Verification:** {row['verification_status']} (evidence tier {row['evidence_tier']} - {row['evidence_tier_label']})")
                st.write(f"**Weakest dimensions:** {row['weakest_dimensions']}")
                st.write(f"**Root cause:** {row['root_cause']}")
                st.write(f"**How to increase the score:** {row['how_to_increase_score']}")
                if "No major quality challenge" not in str(row["quality_challenge_flags"]):
                    st.error(row["quality_challenge_flags"])
                st.caption(f"Scoring trace: {row['score_trace']}")
                st.caption(f"Governance: {row['governance_note']}")

    with tabs[2]:
        qa_status_cols = ["insight_id"] + [c for c in report.columns if c.endswith("_status")]
        st.dataframe(report[qa_status_cols], width="stretch", hide_index=True)
        st.caption("N/A means the check does not apply to that insight, for example a numbers check on a "
                   "claim that quotes no numbers.")
        with st.expander("Full QA notes"):
            note_cols = ["insight_id"] + [c for c in report.columns if c.endswith("_note")]
            st.dataframe(report[note_cols], width="stretch", hide_index=True)

    with tabs[3]:
        evidence_cols = [c for c in ["insight_id", "insight_text", "verification_status", "evidence_tier_label",
                                     "quantitative_metrics_validated", "cross_source_validation", "data_reference"]
                         if c in report.columns]
        st.dataframe(report[evidence_cols], width="stretch", hide_index=True)

    with tabs[4]:
        st.write("What the parser found in each uploaded file:")
        st.dataframe(result["diagnostics"], width="stretch", hide_index=True)
        st.caption(
            "If no comparable metric rows were found, cross-source validation cannot run and decisive claims will "
            "be capped. A readable table needs a header row of column labels and data rows carrying two or more "
            "numbers. Charts pasted as images cannot be read: upload the source table instead."
        )
        st.write("Insight rows as extracted, before scoring:")
        st.dataframe(result["extracted"], width="stretch")

    with tabs[5]:
        st.download_button("Scored report (CSV)", report.to_csv(index=False).encode("utf-8"),
                           "aicf_scored_report.csv", "text/csv", width="stretch")
        st.download_button("Scored report (Excel, multi-sheet)",
                           build_excel_report(report, result["settings"]),
                           "aicf_scored_report.xlsx",
                           "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           width="stretch")
        st.download_button("Validation summary (Word)",
                           build_docx_summary(report, result["settings"]),
                           "aicf_validation_summary.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           width="stretch")
        st.download_button("Insight deck (PowerPoint)",
                           build_pptx_report(report, title="AICF Insight Validation Report",
                                             subtitle="Scored against questionnaire, data and analysis output"),
                           "aicf_scored_insight_deck.pptx",
                           "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                           width="stretch")

with st.expander("How this workflow works"):
    st.write(
        "AICF validates insights that already exist. The questionnaire and the data or analysis output supply the "
        "study context and the numbers. The insight report supplies the claims. Each claim is matched to the "
        "evidence that supports it, scored across the seven weighted dimensions, and given an Insight Confidence "
        "Index out of 100.\n\n"
        "An insight can only reach High Confidence if at least one figure in it was matched to an uploaded source. "
        "Claims that assert a winner, a highest or a lowest are checked against the other options in the same "
        "table or chart, and a claim that names the wrong option is reported as a contradiction rather than "
        "scored down quietly.\n\n"
        "Every scored insight still requires named human researcher sign-off. The index is a governance aid, "
        "not a replacement for researcher judgement."
    )
