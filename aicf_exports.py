"""Export the scored report as a multi-sheet workbook or a Word summary.

v5 offered CSV and PPTX only, while shipping openpyxl and reportlab as unused
dependencies. A 60-column CSV is hard to review, so v6 splits the output into
readable sheets and adds a short Word memo for circulation.
"""

from __future__ import annotations

from io import BytesIO

import pandas as pd

from aicf_config import APP_VERSION, DIMENSIONS

SCORE_COLUMNS = [
    "insight_id", "theme", "insight_text", "ici_score", "ici_classification",
    "verification_status", "evidence_tier", "evidence_tier_label",
    "evidence_strength", "methodological_fit", "triangulation", "interpretability",
    "business_relevance", "actionability", "bias_risk", "weighted_score",
    "weakest_dimensions", "root_cause", "how_to_increase_score", "reliance_level",
]

EVIDENCE_COLUMNS = [
    "insight_id", "insight_text", "evidence_note", "data_reference",
    "quantitative_metrics_validated", "cross_source_validation",
    "quality_challenge_flags", "score_trace", "context_alignment_score",
    "context_alignment_note",
]

GOVERNANCE_COLUMNS = [
    "insight_id", "study_objective", "industry", "technology", "study_type",
    "analytical_technique", "insight_source", "researcher_owner",
    "independent_ai_checker", "benchmark_reference", "governance_note", "scoring_mode",
]


def _subset(report: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    available = [column for column in columns if column in report.columns]
    return report[available].copy() if available else pd.DataFrame()


def build_excel_report(report: pd.DataFrame, settings: dict[str, str]) -> bytes:
    """Multi-sheet workbook: Scores, QA, Evidence trail, Governance, Settings."""
    buffer = BytesIO()
    qa_columns = ["insight_id"] + [c for c in report.columns if c.startswith("qa_")]

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        _subset(report, SCORE_COLUMNS).to_excel(writer, sheet_name="Scores", index=False)
        _subset(report, qa_columns).to_excel(writer, sheet_name="First-layer QA", index=False)
        _subset(report, EVIDENCE_COLUMNS).to_excel(writer, sheet_name="Evidence trail", index=False)
        _subset(report, GOVERNANCE_COLUMNS).to_excel(writer, sheet_name="Governance", index=False)

        weights = pd.DataFrame(
            [{"Dimension": item["label"], "Weight": f"{float(item['weight']):.0%}"} for item in DIMENSIONS.values()]
        )
        settings_frame = pd.DataFrame(
            [{"Setting": key, "Value": str(value)} for key, value in settings.items()]
            + [{"Setting": "Tool version", "Value": APP_VERSION}]
        )
        pd.concat(
            [settings_frame, pd.DataFrame([{"Setting": "", "Value": ""}]), weights.rename(
                columns={"Dimension": "Setting", "Weight": "Value"})],
            ignore_index=True,
        ).to_excel(writer, sheet_name="Run settings", index=False)

        report.to_excel(writer, sheet_name="Full output", index=False)

    return buffer.getvalue()


def build_docx_summary(report: pd.DataFrame, settings: dict[str, str]) -> bytes:
    """A short Word memo a researcher can attach to a sign-off request."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading("AICF Insight Validation Summary", level=0)
    document.add_paragraph(APP_VERSION).runs[0].font.size = Pt(9)

    document.add_heading("Study context", level=1)
    for key, value in settings.items():
        document.add_paragraph(f"{key}: {value}", style="List Bullet")

    document.add_heading("Confidence distribution", level=1)
    if "ici_classification" in report.columns:
        for label, count in report["ici_classification"].value_counts().items():
            document.add_paragraph(f"{label}: {count} insight(s)", style="List Bullet")

    flagged = pd.DataFrame()
    if "quality_challenge_flags" in report.columns:
        flagged = report[~report["quality_challenge_flags"].astype(str).str.contains(
            "No major quality challenge", case=False, na=False)]

    document.add_heading("Insights requiring attention before client use", level=1)
    if flagged.empty:
        document.add_paragraph("No insight was challenged by the validation layer.")
    else:
        for _, row in flagged.head(25).iterrows():
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{row.get('insight_id', '')} - {row.get('ici_classification', '')}").bold = True
            document.add_paragraph(str(row.get("insight_text", ""))[:600])
            document.add_paragraph(f"Challenge: {str(row.get('quality_challenge_flags', ''))[:800]}")
            document.add_paragraph(f"To improve: {str(row.get('how_to_increase_score', ''))[:500]}")

    document.add_heading("Sign-off", level=1)
    document.add_paragraph(
        "Every scored insight requires named human researcher sign-off. The Insight Confidence Index is a "
        "governance aid, not a substitute for researcher judgement."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
